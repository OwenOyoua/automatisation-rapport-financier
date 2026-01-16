"""
Exporteur de rapports au format DOCX (Word).
Génère des documents Word éditables avec mise en forme professionnelle.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

from core.models import BilanFonctionnel, BilanFinancier, PatrimoineEntreprise


class DocxExporter:
    """
    Exporteur pour générer des rapports DOCX professionnels.
    """

    def __init__(self):
        pass

    def export(self, report_data, filename: str, options: Dict[str, Any]) -> str:
        """
        Exporter un rapport au format DOCX.
        
        Args:
            report_data: Données du rapport (BilanFonctionnel, BilanFinancier, etc.)
            filename: Nom du fichier de sortie
            options: Options d'exportation
            
        Returns:
            Chemin du fichier généré
        """
        # Créer le répertoire d'export si nécessaire
        export_dir = Path("exports")
        export_dir.mkdir(exist_ok=True)
        
        file_path = export_dir / filename
        
        # Créer le document Word
        doc = Document()
        
        # Configurer le document
        self.setup_document(doc)
        
        # Contenu du document
        # Page de titre
        self.create_title_page(doc, report_data, options)
        
        # Sommaire
        self.create_table_of_contents(doc)
        
        # Contenu principal selon le type de rapport
        if isinstance(report_data, BilanFonctionnel):
            self.create_bilan_fonctionnel_content(doc, report_data, options)
        elif isinstance(report_data, BilanFinancier):
            self.create_bilan_financier_content(doc, report_data, options)
        elif isinstance(report_data, PatrimoineEntreprise):
            self.create_patrimoine_content(doc, report_data, options)
        
        # Annexes
        if options.get('include_notes', True):
            self.create_annexes(doc, report_data, options)
        
        # Sauvegarder le document
        doc.save(str(file_path))
        
        return str(file_path)

    def setup_document(self, doc: Document):
        """Configurer les paramètres du document."""
        # Configurer les marges
        for section in doc.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)

    def create_title_page(self, doc: Document, report_data, options: Dict[str, Any]):
        """Créer la page de titre."""
        # Ajouter des paragraphes vides pour l'espacement
        for _ in range(8):
            doc.add_paragraph()
        
        # Titre principal
        if isinstance(report_data, BilanFonctionnel):
            title = "BILAN FONCTIONNEL"
        elif isinstance(report_data, BilanFinancier):
            title = "BILAN FINANCIER"
        elif isinstance(report_data, PatrimoineEntreprise):
            title = "PATRIMOINE DE L'ENTREPRISE"
        else:
            title = "RAPPORT FINANCIER"
        
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Espacement
        for _ in range(3):
            doc.add_paragraph()
        
        # Informations sur l'entreprise
        entreprise = options.get('entreprise', 'Entreprise')
        periode = options.get('periode', '2024')
        devise = options.get('devise', 'MAD')
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run("Entreprise: ").bold = True
        info_para.add_run(f"{entreprise}\n")
        info_para.add_run("Période: ").bold = True
        info_para.add_run(f"{periode}\n")
        info_para.add_run("Devise: ").bold = True
        info_para.add_run(f"{devise}\n")
        info_para.add_run("Date d'édition: ").bold = True
        info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y')}")
        
        # Espacement
        for _ in range(3):
            doc.add_paragraph()
        
        # Footer
        if options.get('include_logo', True):
            logo_para = doc.add_paragraph("[Logo de l'entreprise]")
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.runs[0]
            logo_run.font.italic = True
            logo_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Pied de page
        for _ in range(2):
            doc.add_paragraph()
        
        footer_para = doc.add_paragraph("Rapport généré par l'application de comptabilité professionnelle")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

    def create_table_of_contents(self, doc: Document):
        """Créer le sommaire."""
        # Saut de page
        doc.add_page_break()
        
        # Titre du sommaire
        toc_title = doc.add_paragraph("SOMMAIRE")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = toc_title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Espacement
        doc.add_paragraph()
        
        # Contenu du sommaire
        toc_items = [
            ("1. Introduction", "3"),
            ("2. Analyse financière", "4"),
            ("3. Tableaux détaillés", "6"),
            ("4. Ratios et indicateurs", "8"),
            ("5. Recommandations", "10"),
            ("6. Annexes", "12")
        ]
        
        for title, page in toc_items:
            para = doc.add_paragraph()
            para.add_run(title)
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
            para.add_run(f"\t{page}")
        
        doc.add_paragraph()

    def create_bilan_fonctionnel_content(self, doc: Document, bilan: BilanFonctionnel, options: Dict[str, Any]):
        """Créer le contenu du bilan fonctionnel."""
        # Saut de page
        doc.add_page_break()
        
        # Titre
        title = doc.add_paragraph("BILAN FONCTIONNEL")
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Introduction
        intro = doc.add_paragraph()
        intro.add_run("Le bilan fonctionnel présente l'analyse de la structure financière de l'entreprise ")
        intro.add_run("selon l'approche fonctionnelle. Il met en évidence les équilibres fondamentaux : ")
        intro.add_run("Fonds de Roulement Net Global (FRNG), Besoin en Fonds de Roulement (BFR) et Trésorerie.")
        
        doc.add_paragraph()
        
        # Tableau principal
        self.create_bilan_fonctionnel_table(doc, bilan, options)
        
        doc.add_paragraph()
        
        # Analyse
        analysis_title = doc.add_paragraph("ANALYSE FONCTIONNELLE")
        analysis_run = analysis_title.runs[0]
        analysis_run.font.size = Pt(14)
        analysis_run.font.bold = True
        
        analysis_text = self._analyze_bilan_fonctionnel(bilan)
        analysis_para = doc.add_paragraph(analysis_text)
        
        doc.add_paragraph()
        
        # Recommandations
        if options.get('include_ratios', True):
            rec_title = doc.add_paragraph("RECOMMANDATIONS")
            rec_run = rec_title.runs[0]
            rec_run.font.size = Pt(14)
            rec_run.font.bold = True
            
            recommandations = self._get_bilan_fonctionnel_recommendations(bilan)
            for rec in recommandations:
                rec_para = doc.add_paragraph(f"• {rec}")

    def create_bilan_financier_content(self, doc: Document, bilan: BilanFinancier, options: Dict[str, Any]):
        """Créer le contenu du bilan financier."""
        # Saut de page
        doc.add_page_break()
        
        # Titre
        title = doc.add_paragraph("BILAN FINANCIER")
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Introduction
        intro = doc.add_paragraph()
        intro.add_run("Le bilan financier présente la situation patrimoniale de l'entreprise selon ")
        intro.add_run("la présentation comptable classique. Il distingue clairement les actifs ")
        intro.add_run("et les passifs pour évaluer la structure financière et la solvabilité.")
        
        doc.add_paragraph()
        
        # Tableau de l'actif
        actif_title = doc.add_paragraph("ACTIF")
        actif_run = actif_title.runs[0]
        actif_run.font.size = Pt(12)
        actif_run.font.bold = True
        
        self.create_actif_table(doc, bilan, options)
        
        doc.add_paragraph()
        
        # Tableau du passif
        passif_title = doc.add_paragraph("PASSIF")
        passif_run = passif_title.runs[0]
        passif_run.font.size = Pt(12)
        passif_run.font.bold = True
        
        self.create_passif_table(doc, bilan, options)
        
        doc.add_paragraph()
        
        # Analyse
        analysis_title = doc.add_paragraph("ANALYSE FINANCIÈRE")
        analysis_run = analysis_title.runs[0]
        analysis_run.font.size = Pt(14)
        analysis_run.font.bold = True
        
        analysis_text = self._analyze_bilan_financier(bilan)
        analysis_para = doc.add_paragraph(analysis_text)

    def create_patrimoine_content(self, doc: Document, patrimoine: PatrimoineEntreprise, options: Dict[str, Any]):
        """Créer le contenu du patrimoine."""
        # Saut de page
        doc.add_page_break()
        
        # Titre
        title = doc.add_paragraph("PATRIMOINE DE L'ENTREPRISE")
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Introduction
        intro = doc.add_paragraph()
        intro.add_run("L'analyse patrimoniale évalue la valeur réelle de l'entreprise en tenant compte ")
        intro.add_run("de ses actifs économiques et de ses dettes. Elle permet d'apprécier la solidité ")
        intro.add_run("financière et la capacité de l'entreprise à faire face à ses engagements.")
        
        doc.add_paragraph()
        
        # Tableau patrimonial
        self.create_patrimoine_table(doc, patrimoine, options)
        
        doc.add_paragraph()
        
        # Ratios patrimoniaux
        if options.get('include_ratios', True):
            ratios_title = doc.add_paragraph("RATIOS PATRIMONIAUX")
            ratios_run = ratios_title.runs[0]
            ratios_run.font.size = Pt(12)
            ratios_run.font.bold = True
            
            self.create_ratios_table(doc, patrimoine)
            doc.add_paragraph()
        
        # Analyse
        analysis_title = doc.add_paragraph("ANALYSE PATRIMONIALE")
        analysis_run = analysis_title.runs[0]
        analysis_run.font.size = Pt(14)
        analysis_run.font.bold = True
        
        analysis_text = self._analyze_patrimoine(patrimoine)
        analysis_para = doc.add_paragraph(analysis_text)

    def create_bilan_fonctionnel_table(self, doc: Document, bilan: BilanFonctionnel, options: Dict[str, Any]):
        """Créer le tableau du bilan fonctionnel."""
        # Données du tableau
        table_data = [
            ["EMPLOIS ET RESSOURCES", "Montant", "Pourcentage"],
            ["EMPLOIS STABLES", f"{float(bilan.emplois_stables):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Ressources stables", f"{float(bilan.ressources_stables):,.2f} {options.get('devise', 'MAD')}", ""],
            ["FRNG", f"{float(bilan.frng):,.2f} {options.get('devise', 'MAD')}", "100%"],
            ["", "", ""],
            ["ACTIFS CIRCULANTS", f"{float(bilan.actifs_circulants):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Passifs circulants", f"{float(bilan.passifs_circulants):,.2f} {options.get('devise', 'MAD')}", ""],
            ["BFR", f"{float(bilan.bfr):,.2f} {options.get('devise', 'MAD')}", ""],
            ["", "", ""],
            ["TRÉSORERIE ACTIVE", f"{float(bilan.tresorerie_active):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Trésorerie passive", f"{float(bilan.tresorerie_passive):,.2f} {options.get('devise', 'MAD')}", ""],
            ["TRÉSORERIE NETTE", f"{float(bilan.tresorerie_nette):,.2f} {options.get('devise', 'MAD')}", ""],
        ]
        
        # Créer le tableau
        table = doc.add_table(rows=len(table_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Remplir le tableau
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell_data
                
                # Style
                if i == 0:  # En-tête
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif i in [3, 7, 10]:  # Sous-totaux
                    run.font.bold = True
                    self.set_cell_background(cell, RGBColor(230, 230, 230))
                elif i == 11:  # Total
                    run.font.bold = True
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    def create_actif_table(self, doc: Document, bilan: BilanFinancier, options: Dict[str, Any]):
        """Créer le tableau de l'actif."""
        table_data = [
            ["Rubriques", "Montant", "Pourcentage"],
            ["Immobilisations nettes", f"{float(bilan.immobilisations_nettes):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Stocks", f"{float(bilan.stocks):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Créances clients", f"{float(bilan.creances_clients):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Autres créances", f"{float(bilan.autres_creances):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Trésorerie active", f"{float(bilan.tresorerie_active):,.2f} {options.get('devise', 'MAD')}", ""],
            ["TOTAL ACTIF", f"{float(bilan.total_actif):,.2f} {options.get('devise', 'MAD')}", "100%"],
        ]
        
        table = doc.add_table(rows=len(table_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell_data
                
                if i == 0 or i == 6:  # En-tête et total
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    def create_passif_table(self, doc: Document, bilan: BilanFinancier, options: Dict[str, Any]):
        """Créer le tableau du passif."""
        table_data = [
            ["Rubriques", "Montant", "Pourcentage"],
            ["Capital social", f"{float(bilan.capital_social):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Réserves", f"{float(bilan.reserves):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Résultat net", f"{float(bilan.resultat_net):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Capitaux propres", f"{float(bilan.capitaux_propres):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Dettes financières LT", f"{float(bilan.dettes_financieres_lt):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Dettes fournisseurs", f"{float(bilan.dettes_fournisseurs):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Autres dettes CT", f"{float(bilan.autres_dettes_ct):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Trésorerie passive", f"{float(bilan.tresorerie_passive):,.2f} {options.get('devise', 'MAD')}", ""],
            ["TOTAL PASSIF", f"{float(bilan.total_passif):,.2f} {options.get('devise', 'MAD')}", "100%"],
        ]
        
        table = doc.add_table(rows=len(table_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell_data
                
                if i == 0 or i == 9:  # En-tête et total
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif i == 4:  # Sous-total capitaux propres
                    run.font.bold = True
                    self.set_cell_background(cell, RGBColor(230, 230, 230))
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    def create_patrimoine_table(self, doc: Document, patrimoine: PatrimoineEntreprise, options: Dict[str, Any]):
        """Créer le tableau patrimonial."""
        table_data = [
            ["ÉLÉMENTS PATRIMONIAUX", "Montant", "Pourcentage"],
            ["Actifs économiques", f"{float(patrimoine.actifs_economiques):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Dettes financières", f"{float(patrimoine.dettes_financieres):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Actif net comptable", f"{float(patrimoine.actif_net_comptable):,.2f} {options.get('devise', 'MAD')}", ""],
            ["Capitaux propres retraités", f"{float(patrimoine.capitaux_propres_retraites):,.2f} {options.get('devise', 'MAD')}", ""],
            ["PATRIMOINE NET", f"{float(patrimoine.patrimoine_net):,.2f} {options.get('devise', 'MAD')}", "100%"],
        ]
        
        table = doc.add_table(rows=len(table_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell_data
                
                if i == 0 or i == 5:  # En-tête et total
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    def create_ratios_table(self, doc: Document, patrimoine: PatrimoineEntreprise):
        """Créer le tableau des ratios."""
        table_data = [
            ["Ratio", "Valeur", "Interprétation"],
            ["Ratio d'endettement", f"{patrimoine.ratio_endettement or 0:.2%}", self._interpret_ratio(patrimoine.ratio_endettement, 0.5, 0.8)],
            ["Ratio de solvabilité", f"{patrimoine.ratio_solvabilite or 0:.2f}", self._interpret_solvability(patrimoine.ratio_solvabilite)],
            ["Ratio de liquidité", f"{patrimoine.ratio_liquidite or 0:.2f}", self._interpret_ratio(patrimoine.ratio_liquidite, 1.0, 0.8)],
        ]
        
        table = doc.add_table(rows=len(table_data), cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell_data
                
                if i == 0:  # En-tête
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    self.set_cell_background(cell, RGBColor(51, 51, 51))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

    def create_annexes(self, doc: Document, report_data, options: Dict[str, Any]):
        """Créer les annexes du rapport."""
        # Saut de page
        doc.add_page_break()
        
        # Titre
        annexes_title = doc.add_paragraph("ANNEXES")
        title_run = annexes_title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Méthodologie
        methode_title = doc.add_paragraph("MÉTHODOLOGIE")
        methode_run = methode_title.runs[0]
        methode_run.font.size = Pt(14)
        methode_run.font.bold = True
        
        methode_text = f"""
Source des données: Données comptables fournies par l'entreprise.
Période de référence: {options.get('periode', '2024')}
Devise: {options.get('devise', 'MAD')}
Normes comptables: Plan Comptable Marocain
Date d'édition: {datetime.now().strftime('%d/%m/%Y')}

Méthodes de calcul:
• FRNG = Ressources stables - Emplois stables
• BFR = Actifs circulants - Passifs circulants
• Trésorerie nette = Trésorerie active - Trésorerie passive
• Ratio d'endettement = Dettes / Actifs économiques
• Ratio de solvabilité = Capitaux propres / Dettes
• Ratio de liquidité = Actifs liquides / Passifs exigibles
        """
        
        methode_para = doc.add_paragraph(methode_text.strip())
        
        # Notes techniques
        notes_title = doc.add_paragraph("NOTES TECHNIQUES")
        notes_run = notes_title.runs[0]
        notes_run.font.size = Pt(14)
        notes_run.font.bold = True
        
        notes_text = """
Arrondis: Les montants sont arrondis à deux décimales.
Convention de signe: Les soldes positifs indiquent un excédent, les soldes négatifs un déficit.
Pourcentages: Calculés par rapport au total de la rubrique principale.

Avertissement: Ce rapport est basé sur les informations fournies et ne constitue 
pas un conseil en investissement. Une analyse complémentaire peut être nécessaire 
pour une prise de décision éclairée.
        """
        
        notes_para = doc.add_paragraph(notes_text.strip())

    def set_cell_background(self, cell, color):
        """Définir la couleur de fond d'une cellule."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), f"{color.red:02x}{color.green:02x}{color.blue:02x}")
        cell._tc.get_or_add_tcPr().append(shading)

    # Méthodes d'analyse (similaires à celles du PDFExporter)
    def _analyze_bilan_fonctionnel(self, bilan: BilanFonctionnel) -> str:
        """Analyser le bilan fonctionnel."""
        analyses = []
        
        if float(bilan.frng) > 0:
            analyses.append("Le Fonds de Rouillage Net Global est positif, ce qui indique que les ressources stables financent correctement les emplois stables.")
        else:
            analyses.append("Le FRNG est négatif, ce qui révèle une dépendance aux financements à court terme pour couvrir les investissements.")
        
        if float(bilan.bfr) > 0:
            analyses.append("Le Besoin en Fonds de Roulement est positif, ce qui signifie que le cycle d'exploitation nécessite un financement.")
        else:
            analyses.append("Le BFR est négatif, ce qui constitue une ressource de financement issue de l'exploitation.")
        
        if float(bilan.tresorerie_nette) > 0:
            analyses.append("La trésorerie nette est positive, offrant une marge de sécurité financière.")
        elif float(bilan.tresorerie_nette) < 0:
            analyses.append("La trésorerie nette est négative, ce qui peut créer des tensions de liquidité.")
        
        return " ".join(analyses)

    def _analyze_bilan_financier(self, bilan: BilanFinancier) -> str:
        """Analyser le bilan financier."""
        analyses = []
        
        total_actif = float(bilan.total_actif)
        if total_actif > 0:
            ratio_immobilisations = float(bilan.immobilisations_nettes) / total_actif
            ratio_capitaux_propres = float(bilan.capitaux_propres) / total_actif
            
            if ratio_immobilisations > 0.5:
                analyses.append("L'entreprise présente une structure capitalistique marquée avec des immobilisations importantes.")
            
            if ratio_capitaux_propres > 0.5:
                analyses.append("Les capitaux propres constituent la principale source de financement, assurant une bonne autonomie financière.")
            elif ratio_capitaux_propres < 0.2:
                analyses.append("L'entreprise dépend fortement de l'endettement pour son financement.")
        
        return " ".join(analyses)

    def _analyze_patrimoine(self, patrimoine: PatrimoineEntreprise) -> str:
        """Analyser le patrimoine."""
        analyses = []
        
        if patrimoine.ratio_endettement and patrimoine.ratio_endettement < 0.5:
            analyses.append("L'endettement est maîtrisé, offrant une bonne sécurité financière.")
        elif patrimoine.ratio_endettement and patrimoine.ratio_endettement > 0.8:
            analyses.append("L'endettement est élevé et peut compromettre la solvabilité à long terme.")
        
        if patrimoine.ratio_solvabilite and patrimoine.ratio_solvabilite > 1:
            analyses.append("La solvabilité est assurée avec des capitaux propres supérieurs aux dettes.")
        
        return " ".join(analyses)

    def _get_bilan_fonctionnel_recommendations(self, bilan: BilanFonctionnel) -> list:
        """Obtenir les recommandations pour le bilan fonctionnel."""
        recommendations = []
        
        if float(bilan.frng) < 0:
            recommendations.append("Renforcer les ressources stables (augmentation de capital, dettes à long terme)")
        
        if float(bilan.bfr) > float(bilan.frng):
            recommendations.append("Optimiser le cycle d'exploitation pour réduire le BFR")
        
        if float(bilan.tresorerie_nette) < 0:
            recommendations.append("Améliorer la gestion de trésorerie (negociation des délais, cession d'actifs)")
        
        if not recommendations:
            recommendations.append("La structure financière est équilibrée, maintenir la politique actuelle")
        
        return recommendations

    def _interpret_ratio(self, ratio: Optional[float], good_threshold: float, bad_threshold: float) -> str:
        """Interpréter un ratio."""
        if ratio is None:
            return "Non calculable"
        
        if ratio <= good_threshold:
            return "✅ Bon"
        elif ratio >= bad_threshold:
            return "⚠️ À surveiller"
        else:
            return "🟡 Moyen"

    def _interpret_solvability(self, ratio: Optional[float]) -> str:
        """Interpréter le ratio de solvabilité."""
        if ratio is None:
            return "Non calculable"
        
        if ratio > 1:
            return "✅ Solvable"
        elif ratio > 0.5:
            return "🟡 Solvabilité moyenne"
        else:
            return "⚠️ Solvabilité faible"
